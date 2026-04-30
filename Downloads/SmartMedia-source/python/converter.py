import os
import io
import logging
from typing import List, Dict, Union, Optional
from PIL import Image

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger('SmartMediaConverter')

class DocumentConverter:
    def __init__(self):
        pass

    def _ensure_dir(self, path):
        dirname = os.path.dirname(path)
        if dirname and not os.path.exists(dirname):
            os.makedirs(dirname)
    
    def extract_text_from_pdf(self, pdf_path: str) -> Dict:
        """Extract text from PDF file for indexing and search"""
        try:
            from pypdf import PdfReader
        except ImportError:
            return {"success": False, "error": "pypdf not installed", "text": ""}
        
        try:
            reader = PdfReader(pdf_path)
            text = ""
            page_count = len(reader.pages)
            
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"

            # Fallback to PyMuPDF if pypdf extracted little or no text.
            if not text.strip():
                try:
                    import fitz

                    doc = fitz.open(pdf_path)
                    fitz_text = []
                    for page in doc:
                        page_text = page.get_text("text")
                        if page_text:
                            fitz_text.append(page_text)
                    doc.close()
                    text = "\n\n".join(fitz_text)
                    if text.strip():
                        logger.info("PDF text fallback via PyMuPDF succeeded")
                except Exception as fallback_error:
                    logger.warning(f"PDF PyMuPDF fallback failed: {fallback_error}")
            
            return {
                "success": True,
                "text": text.strip(),
                "pages": page_count,
                "file_type": "pdf"
            }
        except Exception as e:
            logger.error(f"PDF text extraction error: {e}")
            return {"success": False, "error": str(e), "text": ""}
    
    def extract_text_from_word(self, docx_path: str) -> Dict:
        """Extract text from Word document"""
        ext = os.path.splitext(docx_path)[1].lower()

        # Legacy .doc files need a different extractor than python-docx.
        if ext == '.doc':
            try:
                import win32com.client  # type: ignore

                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                document = word.Documents.Open(os.path.abspath(docx_path), ReadOnly=1)
                text = document.Content.Text or ""
                document.Close(False)
                word.Quit()
                return {
                    "success": True,
                    "text": text.strip(),
                    "paragraphs": None,
                    "file_type": "doc",
                }
            except ImportError:
                return {"success": False, "error": "python-docx does not support .doc files and pywin32 is not installed", "text": ""}
            except Exception as e:
                logger.error(f"Legacy Word (.doc) extraction error: {e}")
                return {"success": False, "error": str(e), "text": ""}

        try:
            from docx import Document
        except ImportError:
            return {"success": False, "error": "python-docx not installed", "text": ""}
        
        try:
            doc = Document(docx_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join([cell.text for cell in row.cells])
                    text += row_text + "\n"
            
            return {
                "success": True,
                "text": text.strip(),
                "paragraphs": len(doc.paragraphs),
                "file_type": "docx"
            }
        except Exception as e:
            logger.error(f"Word text extraction error: {e}")
            return {"success": False, "error": str(e), "text": ""}
    
    def extract_text_from_txt(self, txt_path: str) -> Dict:
        """Extract text from plain text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            text = None
            
            for encoding in encodings:
                try:
                    with open(txt_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    break  # Success!
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                return {"success": False, "error": "Unable to decode file", "text": ""}
            
            return {
                "success": True,
                "text": text.strip(),
                "file_type": "txt"
            }
        except Exception as e:
            logger.error(f"Text file extraction error: {e}")
            return {"success": False, "error": str(e), "text": ""}
    
    def extract_text_from_document(self, file_path: str) -> Dict:
        """Universal text extraction from any supported document"""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return self.extract_text_from_word(file_path)
        elif ext in ['.txt', '.md', '.log', '.csv']:
            return self.extract_text_from_txt(file_path)
        else:
            return {
                "success": False,
                "error": f"Unsupported file type: {ext}",
                "text": ""
            }

    def pdf_to_word(self, pdf_path: str, output_path: str = None) -> Dict:
        try:
            from pdf2docx import Converter
        except ImportError:
            return {"success": False, "error": "pdf2docx not installed. Run: pip install pdf2docx"}

        try:
            if not output_path:
                output_path = os.path.splitext(pdf_path)[0] + ".docx"

            self._ensure_dir(output_path)
            
            logger.info(f"Converting PDF to Word: {pdf_path} -> {output_path}")
            
            cv = Converter(pdf_path)
            cv.convert(output_path, start=0, end=None)
            cv.close()
            
            return {
                "success": True, 
                "output_path": output_path,
                "message": "Conversion successful"
            }
        except Exception as e:
            logger.error(f"PDF to Word error: {e}")
            return {"success": False, "error": str(e)}

    def word_to_pdf(self, docx_path: str, output_path: str = None) -> Dict:
        try:
            from docx2pdf import convert
        except ImportError:
            return {"success": False, "error": "docx2pdf not installed. Run: pip install docx2pdf"}

        try:
            if not output_path:
                output_path = os.path.splitext(docx_path)[0] + ".pdf"

            self._ensure_dir(output_path)
            
            logger.info(f"Converting Word to PDF: {docx_path} -> {output_path}")
            
            # Note: This requires Microsoft Word installed on Windows
            convert(docx_path, output_path)
            
            return {
                "success": True, 
                "output_path": output_path,
                "message": "Conversion successful"
            }
        except Exception as e:
            logger.error(f"Word to PDF error: {e}")
            return {"success": False, "error": str(e) + " (Ensure Microsoft Word is installed)"}

    def images_to_pdf(self, image_paths: List[str], output_path: str = None) -> Dict:
        try:
            import img2pdf
        except ImportError:
            return {"success": False, "error": "img2pdf not installed. Run: pip install img2pdf"}

        try:
            if not image_paths:
                return {"success": False, "error": "No images provided"}

            if not output_path:
                # Default output name based on first image or timestamp
                dir_name = os.path.dirname(image_paths[0])
                import time
                timestamp = int(time.time())
                output_path = os.path.join(dir_name, f"combined_images_{timestamp}.pdf")

            self._ensure_dir(output_path)
            
            logger.info(f"Converting {len(image_paths)} images to PDF -> {output_path}")
            
            # Write PDF
            with open(output_path, "wb") as f:
                f.write(img2pdf.convert(image_paths))
            
            return {
                "success": True, 
                "output_path": output_path,
                "message": "PDF created successfully"
            }
        except Exception as e:
            logger.error(f"Images to PDF error: {e}")
            return {"success": False, "error": str(e)}

    def merge_pdfs(self, pdf_paths: List[str], output_path: str = None) -> Dict:
        try:
            from pypdf import PdfWriter
        except ImportError:
            return {"success": False, "error": "pypdf not installed. Run: pip install pypdf"}

        try:
            if not pdf_paths:
                return {"success": False, "error": "No PDFs provided"}

            if not output_path:
                dir_name = os.path.dirname(pdf_paths[0])
                import time
                timestamp = int(time.time())
                output_path = os.path.join(dir_name, f"merged_{timestamp}.pdf")

            self._ensure_dir(output_path)
            
            logger.info(f"Merging {len(pdf_paths)} PDFs -> {output_path}")
            
            merger = PdfWriter()
            
            for pdf in pdf_paths:
                merger.append(pdf)
            
            merger.write(output_path)
            merger.close()
            
            return {
                "success": True, 
                "output_path": output_path,
                "message": "PDFs merged successfully"
            }
        except Exception as e:
            logger.error(f"Merge PDFs error: {e}")
            return {"success": False, "error": str(e)}

    def split_pdf(self, pdf_path: str, output_dir: str = None) -> Dict:
        try:
            from pypdf import PdfReader, PdfWriter
        except ImportError:
            return {"success": False, "error": "pypdf not installed. Run: pip install pypdf"}

        try:
            if not output_dir:
                output_dir = os.path.dirname(pdf_path)

            filename = os.path.basename(pdf_path)
            name, ext = os.path.splitext(filename)
            
            logger.info(f"Splitting PDF: {pdf_path}")
            
            reader = PdfReader(pdf_path)
            output_files = []
            
            for i, page in enumerate(reader.pages):
                writer = PdfWriter()
                writer.add_page(page)
                
                output_filename = f"{name}_page_{i+1}{ext}"
                output_path = os.path.join(output_dir, output_filename)
                
                with open(output_path, "wb") as f:
                    writer.write(f)
                
                output_files.append(output_path)
            
            return {
                "success": True, 
                "output_paths": output_files,
                "count": len(output_files),
                "message": f"Split into {len(output_files)} pages"
            }
        except Exception as e:
            logger.error(f"Split PDF error: {e}")
            return {"success": False, "error": str(e)}

    # ===== IMAGE TOOLS =====

    def convert_image(self, image_paths: List[str], output_format: str) -> Dict:
        """Convert images to another format (jpg, png, webp, etc.)"""
        try:
            converted_files = []
            output_format = output_format.lower()
            if output_format == 'jpg': output_format = 'jpeg'
            
            for img_path in image_paths:
                try:
                    img = Image.open(img_path)
                    
                    # Handle RGBA to RGB convert for JPEG
                    if output_format == 'jpeg' and img.mode in ('RGBA', 'P'):
                        img = img.convert('RGB')
                        
                    output_path = os.path.splitext(img_path)[0] + f".{output_format}"
                    
                    # Avoid overwriting original if same extension
                    if output_path == img_path:
                        output_path = os.path.splitext(img_path)[0] + f"_converted.{output_format}"
                        
                    img.save(output_path)
                    converted_files.append(output_path)
                    
                except Exception as e:
                    logger.error(f"Error converting {img_path}: {e}")
            
            return {
                "success": len(converted_files) > 0,
                "output_paths": converted_files,
                "message": f"Converted {len(converted_files)} images"
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def compress_image(self, image_paths: List[str], quality: int = 80) -> Dict:
        """Compress images (JPEG/WebP)"""
        try:
            compressed_files = []
            
            for img_path in image_paths:
                try:
                    img = Image.open(img_path)
                    
                    # Determine format
                    fmt = img.format or 'JPEG'
                    if fmt == 'PNG':
                        # PNG compression is different check optimize flag
                        # Convert to efficient format if desired? No user asked for compress.
                        # Usually Compress means reduce quality for lossy, or optmize for lossless
                        pass
                    
                    output_path = os.path.splitext(img_path)[0] + f"_compressed.{fmt.lower()}"
                    
                    # Handle RGBA to RGB for JPEG
                    if fmt in ('JPEG', 'JPG') and img.mode in ('RGBA', 'P'):
                        img = img.convert('RGB')
                        
                    img.save(output_path, quality=quality, optimize=True)
                    compressed_files.append(output_path)
                    
                except Exception as e:
                    logger.error(f"Error compressing {img_path}: {e}")
                    
            return {
                "success": len(compressed_files) > 0,
                "output_paths": compressed_files,
                "message": f"Compressed {len(compressed_files)} images"
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def resize_image(self, image_paths: List[str], width: int, height: int) -> Dict:
        """Resize images"""
        try:
            resized_files = []
            
            for img_path in image_paths:
                try:
                    img = Image.open(img_path)
                    
                    # Resize
                    img = img.resize((width, height), Image.Resampling.LANCZOS)
                    
                    output_path = os.path.splitext(img_path)[0] + f"_resized.{img.format.lower()}"
                    
                    img.save(output_path)
                    resized_files.append(output_path)
                    
                except Exception as e:
                    logger.error(f"Error resizing {img_path}: {e}")
                    
            return {
                "success": len(resized_files) > 0,
                "output_paths": resized_files,
                "message": f"Resized {len(resized_files)} images"
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def image_to_text(self, image_paths: List[str]) -> Dict:
        """Extract text from images using OCR"""
        try:
            import pytesseract
        except ImportError:
            # Fall back to a simpler approach if pytesseract is not installed
            return {
                "success": False, 
                "error": "pytesseract not installed. Run: pip install pytesseract (and install Tesseract OCR engine)"
            }

        try:
            all_texts = []
            output_path = None
            
            for img_path in image_paths:
                try:
                    img = Image.open(img_path)
                    text = pytesseract.image_to_string(img)
                    all_texts.append({
                        "file": os.path.basename(img_path),
                        "text": text.strip()
                    })
                except Exception as e:
                    logger.error(f"Error extracting text from {img_path}: {e}")
                    all_texts.append({
                        "file": os.path.basename(img_path),
                        "text": "",
                        "error": str(e)
                    })
            
            # Optionally save to a file
            if image_paths:
                dir_name = os.path.dirname(image_paths[0])
                import time
                timestamp = int(time.time())
                output_path = os.path.join(dir_name, f"extracted_text_{timestamp}.txt")
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    for item in all_texts:
                        f.write(f"=== {item['file']} ===\n")
                        f.write(item['text'] + "\n\n")
            
            combined_text = "\n\n".join([item['text'] for item in all_texts if item.get('text')])
            
            return {
                "success": True,
                "texts": all_texts,
                "combined_text": combined_text,
                "output_path": output_path,
                "message": f"Extracted text from {len(all_texts)} images"
            }
        except Exception as e:
            logger.error(f"OCR error: {e}")
            return {"success": False, "error": str(e)}

    def add_watermark(self, image_paths: List[str], text: str, position: str = "bottomRight", opacity: int = 70) -> Dict:
        """Add text watermark to images"""
        try:
            from PIL import ImageDraw, ImageFont
        except ImportError:
            return {"success": False, "error": "Pillow not properly installed"}

        try:
            watermarked_files = []
            
            for img_path in image_paths:
                try:
                    img = Image.open(img_path).convert("RGBA")
                    
                    # Create watermark layer
                    txt_layer = Image.new("RGBA", img.size, (255, 255, 255, 0))
                    draw = ImageDraw.Draw(txt_layer)
                    
                    # Try to get a font, fallback to default
                    try:
                        font = ImageFont.truetype("arial.ttf", size=max(20, img.size[0] // 20))
                    except:
                        font = ImageFont.load_default()
                    
                    # Calculate text size
                    bbox = draw.textbbox((0, 0), text, font=font)
                    text_width = bbox[2] - bbox[0]
                    text_height = bbox[3] - bbox[1]
                    
                    # Position mapping
                    padding = 20
                    positions = {
                        "topLeft": (padding, padding),
                        "topRight": (img.size[0] - text_width - padding, padding),
                        "center": ((img.size[0] - text_width) // 2, (img.size[1] - text_height) // 2),
                        "bottomLeft": (padding, img.size[1] - text_height - padding),
                        "bottomRight": (img.size[0] - text_width - padding, img.size[1] - text_height - padding),
                    }
                    
                    pos = positions.get(position, positions["bottomRight"])
                    
                    # Calculate alpha from opacity percentage
                    alpha = int(255 * opacity / 100)
                    
                    # Draw text with opacity
                    draw.text(pos, text, fill=(255, 255, 255, alpha), font=font)
                    
                    # Composite
                    watermarked = Image.alpha_composite(img, txt_layer)
                    
                    # Save
                    output_path = os.path.splitext(img_path)[0] + "_watermarked.png"
                    watermarked.save(output_path)
                    watermarked_files.append(output_path)
                    
                except Exception as e:
                    logger.error(f"Error watermarking {img_path}: {e}")
            
            return {
                "success": len(watermarked_files) > 0,
                "output_paths": watermarked_files,
                "output_path": watermarked_files[0] if watermarked_files else None,
                "message": f"Watermarked {len(watermarked_files)} images"
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def generate_qr(self, content: str, output_dir: str = None) -> Dict:
        """Generate QR code from text or URL"""
        try:
            import qrcode
        except ImportError:
            return {"success": False, "error": "qrcode not installed. Run: pip install qrcode[pil]"}

        try:
            import time
            
            # Create QR code
            qr = qrcode.QRCode(
                version=1,
                error_correction=qrcode.constants.ERROR_CORRECT_L,
                box_size=10,
                border=4,
            )
            qr.add_data(content)
            qr.make(fit=True)
            
            img = qr.make_image(fill_color="black", back_color="white")
            
            # Determine output path
            if not output_dir:
                output_dir = os.path.expanduser("~/Downloads")
            
            timestamp = int(time.time())
            output_path = os.path.join(output_dir, f"qrcode_{timestamp}.png")
            
            self._ensure_dir(output_path)
            img.save(output_path)
            
            return {
                "success": True,
                "output_path": output_path,
                "message": "QR code generated successfully"
            }
        except Exception as e:
            logger.error(f"QR generation error: {e}")
            return {"success": False, "error": str(e)}

    def apply_color_preset(self, image_paths: List[str], preset: str) -> Dict:
        """Apply color correction presets (warm, cool, vintage, bw, sepia, vivid)"""
        try:
            from PIL import ImageEnhance, ImageFilter
        except ImportError:
            return {"success": False, "error": "Pillow not properly installed"}

        try:
            processed_files = []
            
            for img_path in image_paths:
                try:
                    img = Image.open(img_path)
                    
                    # Ensure RGB mode
                    if img.mode != 'RGB':
                        img = img.convert('RGB')
                    
                    if preset == "warm":
                        # Add warm tones by adjusting color balance
                        r, g, b = img.split()
                        r = r.point(lambda i: min(255, int(i * 1.1)))
                        b = b.point(lambda i: int(i * 0.9))
                        img = Image.merge('RGB', (r, g, b))
                        
                    elif preset == "cool":
                        # Add cool tones
                        r, g, b = img.split()
                        r = r.point(lambda i: int(i * 0.9))
                        b = b.point(lambda i: min(255, int(i * 1.1)))
                        img = Image.merge('RGB', (r, g, b))
                        
                    elif preset == "vintage":
                        # Sepia + reduced saturation + slight vignette effect
                        img = img.convert("L")  # Grayscale first
                        img = img.convert("RGB")
                        r, g, b = img.split()
                        r = r.point(lambda i: min(255, int(i * 1.1)))
                        g = g.point(lambda i: int(i * 0.95))
                        b = b.point(lambda i: int(i * 0.8))
                        img = Image.merge('RGB', (r, g, b))
                        
                    elif preset == "bw":
                        # Black and white
                        img = img.convert("L").convert("RGB")
                        
                    elif preset == "sepia":
                        # Sepia tone
                        gray = img.convert("L")
                        sepia = Image.merge("RGB", (
                            gray.point(lambda x: min(255, x + 50)),
                            gray.point(lambda x: min(255, x + 25)),
                            gray.point(lambda x: max(0, x - 25))
                        ))
                        img = sepia
                        
                    elif preset == "vivid":
                        # Increase saturation and contrast
                        enhancer = ImageEnhance.Color(img)
                        img = enhancer.enhance(1.4)  # 40% more saturation
                        enhancer = ImageEnhance.Contrast(img)
                        img = enhancer.enhance(1.2)  # 20% more contrast
                    
                    # Save
                    ext = os.path.splitext(img_path)[1].lower()
                    if not ext or ext == '.':
                        ext = '.jpg'
                    output_path = os.path.splitext(img_path)[0] + f"_{preset}{ext}"
                    img.save(output_path)
                    processed_files.append(output_path)
                    
                except Exception as e:
                    logger.error(f"Error applying preset to {img_path}: {e}")
            
            return {
                "success": len(processed_files) > 0,
                "output_paths": processed_files,
                "output_path": processed_files[0] if processed_files else None,
                "message": f"Applied {preset} preset to {len(processed_files)} images"
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def remove_background(self, image_paths: List[str]) -> Dict:
        """Remove background from images using rembg"""
        try:
            from rembg import remove
        except ImportError:
            return {"success": False, "error": "rembg not installed. Run: pip install rembg"}

        try:
            processed_files = []
            
            for img_path in image_paths:
                try:
                    with open(img_path, 'rb') as f:
                        input_data = f.read()
                    
                    output_data = remove(input_data)
                    
                    output_path = os.path.splitext(img_path)[0] + "_nobg.png"
                    with open(output_path, 'wb') as f:
                        f.write(output_data)
                    
                    processed_files.append(output_path)
                    
                except Exception as e:
                    logger.error(f"Error removing bg from {img_path}: {e}")
            
            return {
                "success": len(processed_files) > 0,
                "output_paths": processed_files,
                "output_path": processed_files[0] if processed_files else None,
                "message": f"Removed background from {len(processed_files)} images"
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def upscale_image(self, image_paths: List[str], factor: int = 2) -> Dict:
        """Upscale images using simple interpolation (or AI if available)"""
        try:
            upscaled_files = []
            
            for img_path in image_paths:
                try:
                    img = Image.open(img_path)
                    
                    # Calculate new size
                    new_width = img.size[0] * factor
                    new_height = img.size[1] * factor
                    
                    # Upscale using LANCZOS (high quality resampling)
                    upscaled = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                    
                    # Save
                    ext = os.path.splitext(img_path)[1].lower()
                    if not ext or ext == '.':
                        ext = '.png'
                    output_path = os.path.splitext(img_path)[0] + f"_{factor}x{ext}"
                    upscaled.save(output_path)
                    upscaled_files.append(output_path)
                    
                except Exception as e:
                    logger.error(f"Error upscaling {img_path}: {e}")
            
            return {
                "success": len(upscaled_files) > 0,
                "output_paths": upscaled_files,
                "output_path": upscaled_files[0] if upscaled_files else None,
                "message": f"Upscaled {len(upscaled_files)} images to {factor}x"
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

